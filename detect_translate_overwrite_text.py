import io
import os
from google.cloud import vision_v1p3beta1 as vision
from google.cloud import translate_v2 as translate
from PIL import Image, ImageDraw, ImageFont
from PIL import ImageEnhance, ImageFilter
# import cv2
import docx

# Set up the Google Cloud API clients
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = '/path/to/api_key.json'
vision_client = vision.ImageAnnotatorClient()
translate_client = translate.Client()



def load_process_image(image_path):
    # Load image from disk and increase contrast
    image = Image.open(image_path)
    # image = ImageEnhance.Contrast(image).enhance(3)

    # Increase sharpness
    image = ImageEnhance.Sharpness(image).enhance(10)

    # # # Remove noise
    # # image = image.filter(ImageFilter.MedianFilter(3))

    # # # Increase contrast to reduce noise
    # # image = ImageEnhance.Contrast(image).enhance(2)

    # # Save the processed image to disk
    folder_path = os.path.dirname(image_path)
    image.save(os.path.join(folder_path, 'processed_image.png'))

    with io.open(os.path.join(folder_path, 'processed_image.png'), 'rb') as image_file:
    # with io.open(image_path, 'rb') as image_file:
        content = image_file.read()
    input_image = vision.types.Image(content=content)

    # Return the processed image as PIL image
    return input_image

def detect_and_translate_text(image, input_language, target_language):
    """Detects text in an image using the Google Cloud Vision API."""
    # response = vision_client.text_detection(image=image, image_context={"language_hints": [input_language]})
    response = vision_client.text_detection(image=image)
    texts = response.text_annotations[1:]  # Ignore the first element which contains the entire text
    text_blocks = []
    for text in texts:
        vertices = [(vertex.x, vertex.y) for vertex in text.bounding_poly.vertices]
        translated_text = translate_text(text.description, input_language, target_language)      
        text_blocks.append((translated_text, vertices))
    return text_blocks


def translate_text(text, input_language, target_language):
    """Translates text to a target language using the Google Cloud Translate API."""
    # translation = translate_client.translate(text, source_language=input_language, target_language=target_language)
    translation = translate_client.translate(text, target_language=target_language)
    return translation['translatedText']


def text_to_pdf(text_blocks, output_folder):
    """Writes the text blocks to a PDF file."""
    # Create a new Word document
    doc = docx.Document()

    # Concatenate the text blocks of the same line
    text_blocks = concatenate_text_blocks(text_blocks)

    # Sort the text blocks from top to bottom
    # text_blocks = sorted(text_blocks, key=lambda x: x[1][0][1], reverse=True)

    # Write the text sentences to the Word document as new line
    # if word start with capital letter, then it is a new sentence
    for text, vertices in text_blocks:
        if text[0].isupper():
            doc.add_paragraph()
        doc.add_paragraph(text)


    # change the font to Times New Roman size 12
    # style = doc.styles['Normal']
    # font = style.font
    # font.name = 'Times New Roman'
    # font.size = docx.shared.Pt(12)

    # Convert the Word document to PDF
    doc.save(os.path.join(output_folder, 'result.docx'))
    os.system("libreoffice --headless --convert-to pdf result.docx")


# write a method to concatenate the text blocks of the same line together and return a list of sentences
def concatenate_text_blocks(text_blocks):
    # Sort the text blocks from top to bottom
    text_blocks = sorted(text_blocks, key=lambda x: x[1][0][1], reverse=True)
    # Concatenate the text blocks of the same line
    text_blocks_concatenated = []
    for text, vertices in text_blocks:
        if len(text_blocks_concatenated) == 0:
            text_blocks_concatenated.append((text, vertices))
        else:
            if vertices[0][1] - text_blocks_concatenated[-1][1][0][1] < 50:
                text_blocks_concatenated[-1] = (text_blocks_concatenated[-1][0] + ' ' + text, text_blocks_concatenated[-1][1])
            else:
                text_blocks_concatenated.append((text, vertices))
    return text_blocks_concatenated

def draw_text(input_pil_image, text_blocks, font, font_size=20, color='black', color_bg='black'):
    # Create a new PIL Image object with the same size and mode as the input image
    output_pil_image = Image.new(input_pil_image.mode, input_pil_image.size)

    # Copy the input image to the output image
    output_pil_image.paste(input_pil_image)

    # Draw the text on the output image
    draw = ImageDraw.Draw(output_pil_image)
    # font = ImageFont.truetype(font, font_size)
    font = ImageFont.load_default()
    for text, vertices in text_blocks:
        for t, v in zip(text.split('\n'), vertices):
            x0, y0 = v
            x1, y1 = vertices[2]
            x0, x1 = sorted([x0, x1])
            y0, y1 = sorted([y0, y1])
            # increase the size of the box by 30% to avoid cropping letters
            old_x0, old_y0, old_x1, old_y1 = x0 - 0.3 * (x1 - x0), y0 - 0.3 * (y1 - y0), x1 + 0.3 * (x1 - x0), y1 + 0.3 * (y1 - y0)
            x0, y0, x1, y1 = x0 - 0.3 * (x1 - x0), y0 - 0.3 * (y1 - y0), x1 + 0.3 * (x1 - x0), y1 + 0.3 * (y1 - y0)
            region = output_pil_image.crop((x0, y0, x1, y1))
            colors = region.getcolors(region.size[0] * region.size[1])
            # choose the most common color except those that are too close to black or white
            background_color = max(colors, key=lambda x: x[0])[1]
            # print(colors)
            while sum(background_color) < 10 or sum(background_color) > 850:
                x0, y0, x1, y1 = x0 - 0.05 * (x1 - x0), y0 - 0.05 * (y1 - y0), x1 + 0.05 * (x1 - x0), y1 + 0.05 * (y1 - y0)
                region = output_pil_image.crop((x0, y0, x1, y1))
                colors = region.getcolors(region.size[0] * region.size[1])
                # choose second most common color except those that are too close to black or white
                if sum(max(colors, key=lambda x: x[0])[1]) < 10 or sum(max(colors, key=lambda x: x[0])[1]) > 850:
                    background_color = sorted(colors, key=lambda x: x[0])[-2][1]
                else:
                    background_color = max(colors, key=lambda x: x[0])[1]
                # print(colors)

               
            # restore the size of the box
            draw.rectangle([old_x0, old_y0, old_x1, old_y1], fill=background_color)
            # increase the size of the font by 30%
            # font = ImageFont.truetype(font, int(1.3 * font_size))
            draw.text(v, t, font=font, fill=color)

        # print(vertices)
        # draw.rectangle(vertices, outline=color)
        # draw_text(output_image, text, vertices[0], font_size, color)

    return output_pil_image


# Load the input image
# Load the input image
input_image_path = 'YOUR_IMAGE_PATH'
output_folder = 'YOUR_OUTPUT_FOLDER'
output_image_path = os.path.join(output_folder, 'result.jpg')

input_image = load_process_image(input_image_path)



# The language of the text is chinese
input_language = 'zh'
# Translate the text to Korean
# target_language = 'ko'
target_language = 'en'

# Detect text in the input image using the Google Cloud Vision API
text_blocks = detect_and_translate_text(input_image, input_language, target_language)


# write a method to concatenate the text blocks of the same line
text_to_pdf(text_blocks, output_folder)

# Draw the translated text at the top of the input image
input_pil_image = Image.open(input_image_path)

font = '/usr/share/fonts/truetype/sinhala/lklug.ttf'
output_image = draw_text(input_pil_image, text_blocks, font)


# Save the output image
output_image.save(output_image_path)