import pytesseract
import cv2
import docx
import numpy as np
from PIL import Image
import os
import openai
import re 


openai.api_key = "YOUR_API_KEY"

GPT3_MODELS = {
    'davinci': 'text-davinci-003',
    'curie': 'text-curie-001',
    'babbage': 'text-babbage-001',
    'ada': 'text-ada-001'
}

CODEX_MODELS = {
    'davinci': 'code-davinci-002',
    'cushman': 'code-cushman-001'
}


def grammar_checker(prompt, model=GPT3_MODELS['davinci'], temperature=0.0, max_tokens=1000, top_p=1.0, frequency_penalty=0.0, presence_penalty=0.0):
        response = openai.Completion.create(
            model=model,
            prompt=prompt,
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=top_p,
            frequency_penalty=frequency_penalty,
            presence_penalty=presence_penalty,
            n=1
        )
        return response['choices'][0]['text'].strip()

# Path of working folder on Disk
# Read jpeg image to process


folder_path = "/media/full/DATA/Software/Programming_fun/Extract_text_from_image"
# check if the image exists or not
if not os.path.exists(os.path.join(folder_path,"test.jpg")):
    print("Image not found")
    exit()
img = cv2.imread(os.path.join(folder_path,"test.jpg"))

# Convert to gray
img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

# Increase contrast to reduce noise
img = cv2.convertScaleAbs(img, alpha=1.5, beta=0)

# # Apply blur to smooth out the edges
# img = cv2.GaussianBlur(img, (5, 5), 0)



# Write image after removed noise
temp = os.path.join(folder_path,"test_removed.jpeg")
cv2.imwrite(temp, img)

#  Apply threshold to get image with only black and white
img = cv2.adaptiveThreshold(img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)

# Write the image after apply opencv to do some ...
cv2.imwrite(os.path.join(folder_path,"test_thres.png"), img)

# Recognize text with tesseract for python
result = pytesseract.image_to_string(Image.open("test_thres.png"))

# Remove template files
# os.remove(temp)
# os.remove(os.path.join(folder_path,"test_thres.png"))

# Filter result to remove not Unicode or ASCII characters
result1 = "".join([c if ord(c) < 128 else "" for c in result]).strip()

# result = re.sub('[\n\f\t\r]+', ' ', result).replace('  ', ' ')
doc = docx.Document()

# change the font to Times New Roman size 12
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = docx.shared.Pt(12)

doc.add_paragraph(result1)
doc.save(os.path.join(folder_path,'result.docx'))
os.system("libreoffice --headless --convert-to pdf result.docx")

# Improve the text with GPT-3
formatted_text = grammar_checker(result)

# Save the text in a docx file
doc = docx.Document()
doc.add_paragraph(formatted_text)
doc.save(os.path.join(folder_path,'result_AI.docx'))

#convert docx to pdf
os.system("libreoffice --headless --convert-to pdf result_AI.docx")

